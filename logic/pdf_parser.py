"""
PDF and Word document parsing with positional metadata for ProtoScore V2.

Extracts text blocks and tables from protocol documents while preserving
page numbers, bounding boxes, section hierarchy, and structural context.
Uses pdfplumber for table extraction and PyMuPDF for text with coordinates.
"""

from dataclasses import dataclass, field
from typing import Optional
import io
import re
import statistics

import fitz  # PyMuPDF
import pdfplumber


# ---------------------------------------------------------------------------
# Data Models
# ---------------------------------------------------------------------------

@dataclass
class TextBlock:
    """A block of text with positional metadata."""
    text: str
    page_number: int                # 1-indexed
    bbox: tuple[float, float, float, float]  # (x0, y0, x1, y1)
    block_type: str = "paragraph"   # "heading", "paragraph", "list_item"
    section_path: list[str] = field(default_factory=list)
    font_size: float = 0.0
    is_bold: bool = False


@dataclass
class TableCell:
    """A single cell in a table."""
    text: str
    row: int
    col: int
    bbox: Optional[tuple[float, float, float, float]] = None


@dataclass
class ExtractedTable:
    """A table extracted from a document page."""
    cells: list[TableCell] = field(default_factory=list)
    headers: list[str] = field(default_factory=list)
    rows: list[list[str]] = field(default_factory=list)
    page_number: int = 0
    bbox: Optional[tuple[float, float, float, float]] = None
    caption: Optional[str] = None
    section_path: list[str] = field(default_factory=list)


@dataclass
class ParsedPage:
    """All extracted content from a single page."""
    page_number: int
    text_blocks: list[TextBlock] = field(default_factory=list)
    tables: list[ExtractedTable] = field(default_factory=list)
    raw_text: str = ""


@dataclass
class ParsedDocument:
    """Complete parsed document with all pages."""
    pages: list[ParsedPage] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)
    total_pages: int = 0
    section_hierarchy: list[dict] = field(default_factory=list)

    def get_full_text(self) -> str:
        """Concatenate all page text for overview."""
        return "\n\n".join(p.raw_text for p in self.pages if p.raw_text)

    def get_section_text(self, section_name: str) -> str:
        """Get all text from blocks matching a section path."""
        texts = []
        for page in self.pages:
            for block in page.text_blocks:
                if any(section_name.lower() in s.lower() for s in block.section_path):
                    texts.append(block.text)
        return "\n".join(texts)

    def get_tables_in_section(self, section_name: str) -> list[ExtractedTable]:
        """Get all tables from pages matching a section."""
        tables = []
        for page in self.pages:
            for table in page.tables:
                if any(section_name.lower() in s.lower() for s in table.section_path):
                    tables.append(table)
        return tables


# ---------------------------------------------------------------------------
# Section heading detection
# ---------------------------------------------------------------------------

# Common section number patterns in clinical protocols
HEADING_PATTERN = re.compile(
    r'^(?:'
    r'(?:\d+\.?\s)|'                           # "6 " or "6. "
    r'(?:\d+\.\d+\.?\s)|'                      # "6.1 " or "6.1. "
    r'(?:\d+\.\d+\.\d+\.?\s)|'                 # "6.1.1 "
    r'(?:SECTION\s+\d+)|'                       # "SECTION 6"
    r'(?:APPENDIX\s+[A-Z])|'                    # "APPENDIX A"
    r'(?:TABLE\s+\d+)|'                         # "TABLE 1"
    r'(?:FIGURE\s+\d+)'                         # "FIGURE 1"
    r')',
    re.IGNORECASE
)


def _is_heading(text: str, font_size: float, is_bold: bool, median_size: float) -> bool:
    """Determine if a text block is a section heading."""
    text_stripped = text.strip()
    if not text_stripped or len(text_stripped) > 200:
        return False

    # Numbered section pattern
    if HEADING_PATTERN.match(text_stripped):
        return True

    # Large/bold text relative to page median
    if font_size > 0 and median_size > 0:
        if font_size >= median_size + 2 or (is_bold and font_size >= median_size):
            # Headings are usually short
            if len(text_stripped) < 120:
                return True

    # ALL CAPS lines that are short (common in protocols)
    if text_stripped.isupper() and 3 < len(text_stripped) < 80:
        return True

    return False


def _point_in_rect(x: float, y: float, rect: tuple) -> bool:
    """Check if a point falls within a rectangle."""
    return rect[0] <= x <= rect[2] and rect[1] <= y <= rect[3]


def _block_overlaps_table(block_bbox: tuple, table_bboxes: list[tuple]) -> bool:
    """Check if a text block overlaps with any table bounding box."""
    bx0, by0, bx1, by1 = block_bbox
    cx, cy = (bx0 + bx1) / 2, (by0 + by1) / 2  # center point

    for tbbox in table_bboxes:
        if _point_in_rect(cx, cy, tbbox):
            return True
    return False


# ---------------------------------------------------------------------------
# PDF Parsing
# ---------------------------------------------------------------------------

def parse_protocol_pdf(file_bytes: bytes) -> ParsedDocument:
    """
    Parse a protocol PDF into a structured ParsedDocument.

    Uses PyMuPDF for text extraction with font/position metadata and
    pdfplumber for table extraction with bounding boxes.

    Args:
        file_bytes: Raw PDF file bytes

    Returns:
        ParsedDocument with pages, text blocks, tables, and section hierarchy
    """
    doc = ParsedDocument()

    # Open with both libraries
    fitz_doc = fitz.open(stream=file_bytes, filetype="pdf")
    plumber_pdf = pdfplumber.open(io.BytesIO(file_bytes))

    doc.total_pages = len(fitz_doc)
    doc.metadata = dict(fitz_doc.metadata) if fitz_doc.metadata else {}

    # Track current section path for context
    current_sections: list[str] = []

    for page_idx in range(len(fitz_doc)):
        page_num = page_idx + 1
        fitz_page = fitz_doc[page_idx]
        plumber_page = plumber_pdf.pages[page_idx] if page_idx < len(plumber_pdf.pages) else None

        parsed_page = ParsedPage(page_number=page_num)
        parsed_page.raw_text = fitz_page.get_text("text")

        # --- Extract tables via pdfplumber ---
        table_bboxes = []
        if plumber_page:
            plumber_tables = plumber_page.find_tables()
            for pt in plumber_tables:
                bbox = (pt.bbox[0], pt.bbox[1], pt.bbox[2], pt.bbox[3])
                table_bboxes.append(bbox)

                rows_data = pt.extract()
                if not rows_data:
                    continue

                ext_table = ExtractedTable(
                    page_number=page_num,
                    bbox=bbox,
                    section_path=list(current_sections),
                )

                # First row as headers
                ext_table.headers = [str(c) if c else "" for c in rows_data[0]]
                ext_table.rows = []

                for r_idx, row in enumerate(rows_data):
                    row_strs = [str(c) if c else "" for c in row]
                    if r_idx > 0:
                        ext_table.rows.append(row_strs)
                    for c_idx, cell_text in enumerate(row_strs):
                        ext_table.cells.append(TableCell(
                            text=cell_text,
                            row=r_idx,
                            col=c_idx,
                        ))

                parsed_page.tables.append(ext_table)

        # --- Extract text blocks via PyMuPDF ---
        text_dict = fitz_page.get_text("dict")
        font_sizes = []

        # First pass: collect font sizes for median calculation
        for block in text_dict.get("blocks", []):
            if block.get("type") != 0:  # text blocks only
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    if span.get("text", "").strip():
                        font_sizes.append(span["size"])

        median_size = statistics.median(font_sizes) if font_sizes else 12.0

        # Second pass: extract text blocks with metadata
        for block in text_dict.get("blocks", []):
            if block.get("type") != 0:
                continue

            block_bbox = (block["bbox"][0], block["bbox"][1],
                          block["bbox"][2], block["bbox"][3])

            # Skip text blocks that overlap with tables
            if _block_overlaps_table(block_bbox, table_bboxes):
                continue

            # Combine all spans in this block
            block_text_parts = []
            max_font_size = 0.0
            any_bold = False

            for line in block.get("lines", []):
                line_text = ""
                for span in line.get("spans", []):
                    line_text += span.get("text", "")
                    size = span.get("size", 0)
                    if size > max_font_size:
                        max_font_size = size
                    flags = span.get("flags", 0)
                    if flags & 2**4:  # bold flag
                        any_bold = True
                block_text_parts.append(line_text)

            full_text = "\n".join(block_text_parts).strip()
            if not full_text:
                continue

            # Classify block type
            is_heading = _is_heading(full_text, max_font_size, any_bold, median_size)
            block_type = "heading" if is_heading else "paragraph"

            # Update section tracking
            if is_heading:
                # Determine nesting level from numbering
                heading_text = full_text.strip()
                dot_count = heading_text.split()[0].count('.') if heading_text.split() else 0

                if dot_count == 0:
                    current_sections = [heading_text]
                elif dot_count == 1:
                    current_sections = current_sections[:1] + [heading_text]
                else:
                    current_sections = current_sections[:2] + [heading_text]

            text_block = TextBlock(
                text=full_text,
                page_number=page_num,
                bbox=block_bbox,
                block_type=block_type,
                section_path=list(current_sections),
                font_size=max_font_size,
                is_bold=any_bold,
            )
            parsed_page.text_blocks.append(text_block)

        doc.pages.append(parsed_page)

    # Build section hierarchy
    doc.section_hierarchy = _build_section_hierarchy(doc)

    fitz_doc.close()
    plumber_pdf.close()

    return doc


def _build_section_hierarchy(doc: ParsedDocument) -> list[dict]:
    """Build a flat list of section entries with page numbers."""
    sections = []
    seen = set()

    for page in doc.pages:
        for block in page.text_blocks:
            if block.block_type == "heading":
                key = block.text.strip()[:100]
                if key not in seen:
                    seen.add(key)
                    sections.append({
                        "title": block.text.strip(),
                        "page_number": block.page_number,
                        "section_path": block.section_path,
                    })

    return sections


# ---------------------------------------------------------------------------
# Word Document Parsing
# ---------------------------------------------------------------------------

def parse_protocol_docx(file_bytes: bytes) -> ParsedDocument:
    """
    Parse a Word document (.docx) into a ParsedDocument.

    Uses python-docx for text and table extraction. Note: Word documents
    don't have page-level positioning, so bboxes will be None and page
    numbers are estimated from paragraph count.

    Args:
        file_bytes: Raw .docx file bytes

    Returns:
        ParsedDocument with estimated page assignments
    """
    from docx import Document as DocxDocument

    docx_doc = DocxDocument(io.BytesIO(file_bytes))
    doc = ParsedDocument()
    doc.metadata = {
        "title": docx_doc.core_properties.title or "",
        "author": docx_doc.core_properties.author or "",
    }

    # Estimate pages: ~40 paragraphs per page for protocol documents
    PARAS_PER_PAGE = 40
    current_sections: list[str] = []
    current_page_blocks: list[TextBlock] = []
    current_page_tables: list[ExtractedTable] = []
    page_num = 1
    para_count = 0

    for para in docx_doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        para_count += 1

        # Estimate page breaks
        if para_count > PARAS_PER_PAGE:
            doc.pages.append(ParsedPage(
                page_number=page_num,
                text_blocks=current_page_blocks,
                tables=current_page_tables,
                raw_text="\n".join(b.text for b in current_page_blocks),
            ))
            page_num += 1
            current_page_blocks = []
            current_page_tables = []
            para_count = 1

        # Detect headings from Word styles
        is_heading = para.style.name.startswith("Heading") if para.style else False
        if not is_heading:
            is_heading = _is_heading(text, 0, False, 0)

        if is_heading:
            current_sections = [text]

        block = TextBlock(
            text=text,
            page_number=page_num,
            bbox=(0, 0, 0, 0),  # No position data in Word
            block_type="heading" if is_heading else "paragraph",
            section_path=list(current_sections),
            is_bold=any(run.bold for run in para.runs if run.bold is not None),
        )
        current_page_blocks.append(block)

    # Extract tables
    for table in docx_doc.tables:
        ext_table = ExtractedTable(page_number=page_num, section_path=list(current_sections))
        for r_idx, row in enumerate(table.rows):
            row_texts = [cell.text.strip() for cell in row.cells]
            if r_idx == 0:
                ext_table.headers = row_texts
            else:
                ext_table.rows.append(row_texts)
            for c_idx, cell_text in enumerate(row_texts):
                ext_table.cells.append(TableCell(text=cell_text, row=r_idx, col=c_idx))
        current_page_tables.append(ext_table)

    # Flush last page
    if current_page_blocks or current_page_tables:
        doc.pages.append(ParsedPage(
            page_number=page_num,
            text_blocks=current_page_blocks,
            tables=current_page_tables,
            raw_text="\n".join(b.text for b in current_page_blocks),
        ))

    doc.total_pages = page_num
    doc.section_hierarchy = _build_section_hierarchy(doc)

    return doc


# ---------------------------------------------------------------------------
# Page Rendering
# ---------------------------------------------------------------------------

def render_page_with_highlight(
    file_bytes: bytes,
    page_number: int,
    bbox: Optional[tuple[float, float, float, float]] = None,
    margin: int = 50,
) -> bytes:
    """
    Render a PDF page as a PNG image with optional highlight.

    Args:
        file_bytes: Raw PDF bytes
        page_number: 1-indexed page number
        bbox: Optional bounding box to highlight
        margin: Extra margin around highlight (unused, for future crop support)

    Returns:
        PNG image bytes
    """
    doc = fitz.open(stream=file_bytes, filetype="pdf")

    if page_number < 1 or page_number > len(doc):
        doc.close()
        return b""

    page = doc[page_number - 1]

    if bbox:
        rect = fitz.Rect(bbox)
        annot = page.add_highlight_annot(rect)
        annot.set_colors(stroke=(0, 0.64, 0.88))  # ProtoScore teal
        annot.set_opacity(0.4)
        annot.update()

    # Render at 2x for readability
    pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
    img_bytes = pix.tobytes("png")

    doc.close()
    return img_bytes
